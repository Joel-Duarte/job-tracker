from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.main import app
from app.schemas.candidate_profile import CVAnonymizationResult
from app.services.matcher import compute_programmatic_skill_match


def test_programmatic_skill_matcher_aliases_and_ratios():
    candidate_skills = [
        "Python",
        "PostgreSQL",
        "Kubernetes",
        "Docker",
        "FastAPI",
        "Kafka",
    ]
    jd_text = (
        "We are looking for a Senior Engineer with deep knowledge of Python, Postgres, "
        "and k8s container orchestration. Must build high-throughput FastAPI backend services "
        "and distributed Kafka event streaming pipelines. Experience with Docker is a plus."
    )

    result = compute_programmatic_skill_match(candidate_skills, jd_text)
    assert result["candidate_total_skills"] == 6
    assert result["programmatic_score"] == 100
    assert "Python" in result["matching_skills"]
    assert "PostgreSQL" in result["matching_skills"]
    assert "Kubernetes" in result["matching_skills"]
    assert "Docker" in result["matching_skills"]
    assert "FastAPI" in result["matching_skills"]
    assert "Kafka" in result["matching_skills"]


def test_programmatic_skill_matcher_coverage_ratio_50_percent():
    # User has 10 skills, JD requires 20 skills. User matches 10 of them -> score is exactly 50%
    candidate_skills = [
        "Python",
        "FastAPI",
        "PostgreSQL",
        "Docker",
        "Kubernetes",
        "Redis",
        "AWS",
        "Git",
        "CI/CD",
        "TypeScript",
    ]
    jd_required_skills = [
        # 10 matching
        "Python",
        "FastAPI",
        "PostgreSQL",
        "Docker",
        "Kubernetes",
        "Redis",
        "AWS",
        "Git",
        "CI/CD",
        "TypeScript",
        # 10 missing
        "Rust",
        "Go",
        "GraphQL",
        "Apache Spark",
        "Terraform",
        "Solidity",
        "React",
        "Next.js",
        "TailwindCSS",
        "C++",
    ]
    result = compute_programmatic_skill_match(
        candidate_skills=candidate_skills,
        jd_text="Engineering role requirements...",
        jd_required_skills=jd_required_skills,
    )
    assert result["matched_count"] == 10
    assert result["total_required_count"] == 20
    assert result["programmatic_score"] == 50
    assert len(result["matching_skills"]) == 10
    assert len(result["missing_skills"]) == 10
    assert "Rust" in result["missing_skills"]
    assert "Python" in result["matching_skills"]


def test_programmatic_skill_matcher_zero_skills_edge_case():
    candidate_skills = ["Python", "FastAPI"]
    result = compute_programmatic_skill_match(
        candidate_skills=candidate_skills,
        jd_text="General role without specific technical tools mentioned.",
        jd_required_skills=[],
    )
    assert result["programmatic_score"] is None
    assert result["matched_count"] == 0
    assert result["total_required_count"] == 0


@pytest.mark.asyncio
async def test_candidate_profile_crud_and_anonymization(db_session: AsyncSession):
    app.dependency_overrides[get_db] = lambda: db_session
    raw_cv = (
        "John Doe, 123 Main St, San Francisco CA\n"
        "Staff Engineer at Stripe (2018-2022)\n"
        "Built distributed billing pipelines using Python, FastAPI, PostgreSQL, and AWS."
    )

    from app.schemas.candidate_profile import DomainExperienceItem

    mock_anonymized = CVAnonymizationResult(
        anonymized_resume=(
            "Staff Engineer at [Fintech Enterprise] (4 years)\n"
            "Built distributed billing pipelines using Python, FastAPI, PostgreSQL, and AWS."
        ),
        extracted_skills=[
            "Python",
            "FastAPI",
            "PostgreSQL",
            "AWS",
            "Distributed Systems",
        ],
        total_years_experience=6.0,
        domain_expertise=["Fintech", "Distributed Systems"],
        domain_breakdown=[
            DomainExperienceItem(
                domain="Distributed Systems", years=5.0, is_active=True
            ),
            DomainExperienceItem(domain="Fintech", years=3.5, is_active=True),
        ],
        spoken_languages=[
            {"language": "English", "proficiency": "Native"},
            {"language": "Spanish", "proficiency": "B2"},
        ],
        summary="Experienced Staff Engineer with fintech and distributed systems expertise.",
    )

    with patch(
        "app.services.evaluation_worker.anonymize_and_parse_cv", new_callable=AsyncMock
    ) as mock_anonymize:
        mock_anonymize.return_value = mock_anonymized

        app.dependency_overrides[get_db] = lambda: db_session
        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as client:
            # 1. Enqueue CV Task
            resp = await client.post("/api/v1/profile/cv", json={"raw_text": raw_cv})
            assert resp.status_code == 202
            data = resp.json()
            task_id = data["task_id"]
            assert data["status"] in ["QUEUED", "PROCESSING"]

            # Process task in test session
            from app.services.evaluation_worker import process_evaluation_task

            await process_evaluation_task(task_id, db=db_session)

            # 2. Check Task Status
            task_resp = await client.get(f"/api/v1/profile/cv/tasks/{task_id}")
            assert task_resp.status_code == 200
            task_data = task_resp.json()
            assert task_data["status"] == "COMPLETED"
            assert task_data["stage"] == "COMPLETE"

            # 3. Get Active CV Profile
            get_resp = await client.get("/api/v1/profile/cv")
            assert get_resp.status_code == 200
            active_data = get_resp.json()
            assert active_data is not None
            assert "Python" in active_data["extracted_skills"]
            assert active_data["years_of_experience"] == 6.0
            assert len(active_data["domain_experience"]) == 2
            assert (
                active_data["domain_experience"][0]["domain"] == "Distributed Systems"
            )
            assert active_data["domain_experience"][0]["years"] == 5.0
            assert active_data["domain_experience"][0]["is_active"] is True
            assert "[Fintech Enterprise]" in active_data["anonymized_text"]
            assert len(active_data["spoken_languages"]) == 2
            assert active_data["spoken_languages"][0]["language"] == "English"
            assert active_data["spoken_languages"][0]["proficiency"] == "Native"

            # 4. Patch CV (years_of_experience, domain_experience, spoken_languages)
            patch_resp = await client.patch(
                f"/api/v1/profile/cv/{active_data['id']}",
                json={
                    "years_of_experience": 7.0,
                    "anonymized_text": "Updated Custom Sanitized CV",
                    "spoken_languages": [
                        {"language": "English", "proficiency": "Native"},
                        {"language": "German", "proficiency": "C1"},
                    ],
                    "domain_experience": [
                        {
                            "domain": "Distributed Systems",
                            "years": 5.5,
                            "is_active": True,
                        },
                        {"domain": "Fintech", "years": 3.5, "is_active": False},
                        {"domain": "Cloud & DevOps", "years": 2.0, "is_active": True},
                    ],
                },
            )
            assert patch_resp.status_code == 200
            patched_data = patch_resp.json()
            assert patched_data["years_of_experience"] == 7.0
            assert len(patched_data["domain_experience"]) == 3
            assert patched_data["domain_experience"][1]["is_active"] is False
            assert "Cloud & DevOps" in patched_data["domain_expertise"]
            assert len(patched_data["spoken_languages"]) == 2
            assert patched_data["spoken_languages"][1]["language"] == "German"
            assert patched_data["spoken_languages"][1]["proficiency"] == "C1"

            # 5. Delete CV
            del_resp = await client.delete(f"/api/v1/profile/cv/{active_data['id']}")
            assert del_resp.status_code == 204

            # 6. Verify Active CV is None
            get_after_del = await client.get("/api/v1/profile/cv")
            assert get_after_del.status_code == 200
            assert get_after_del.json() is None

    app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_parse_cv_document_file_endpoint():
    import io

    import docx

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as client:
        # 1. Test .txt file upload
        txt_content = (
            b"John Candidate - Full Stack Engineer with Vue and Python experience."
        )
        resp_txt = await client.post(
            "/api/v1/profile/cv/parse-file",
            files={"file": ("resume.txt", txt_content, "text/plain")},
        )
        assert resp_txt.status_code == 200
        data_txt = resp_txt.json()
        assert "John Candidate" in data_txt["text"]
        assert data_txt["filename"] == "resume.txt"

        # 2. Test .docx file upload
        doc = docx.Document()
        doc.add_paragraph("Jane Developer - Lead Software Architect")
        doc.add_paragraph("Expert in FastAPI, Vue, and PostgreSQL")
        docx_buf = io.BytesIO()
        doc.save(docx_buf)

        resp_docx = await client.post(
            "/api/v1/profile/cv/parse-file",
            files={
                "file": (
                    "jane_cv.docx",
                    docx_buf.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert resp_docx.status_code == 200
        data_docx = resp_docx.json()
        assert "Jane Developer" in data_docx["text"]
        assert "FastAPI, Vue, and PostgreSQL" in data_docx["text"]

        # 3. Test empty file upload validation
        resp_empty = await client.post(
            "/api/v1/profile/cv/parse-file",
            files={"file": ("empty.txt", b"", "text/plain")},
        )
        assert resp_empty.status_code == 400
        assert "Uploaded file is empty" in resp_empty.json()["detail"]


def test_spoken_language_match_models_and_schemas():
    from app.schemas.llm import (
        ExtractedJobSpec,
        JobAssessmentResult,
        LanguageMatchResult,
        SpokenLanguageRequirement,
    )

    spec = ExtractedJobSpec(
        position="Senior Backend Engineer",
        company="Tech Corp",
        detected_language="German",
        required_spoken_languages=[
            SpokenLanguageRequirement(
                language="German", requirement="mandatory", proficiency="Fluent / C1"
            )
        ],
    )
    assert spec.detected_language == "German"
    assert len(spec.required_spoken_languages) == 1
    assert spec.required_spoken_languages[0].requirement == "mandatory"

    assessment = JobAssessmentResult(
        position="Senior Backend Engineer",
        company="Tech Corp",
        fit_score=92,
        language_match=LanguageMatchResult(
            is_matched=False,
            detected_jd_language="German",
            required_languages=[
                SpokenLanguageRequirement(
                    language="German",
                    requirement="mandatory",
                    proficiency="Fluent / C1",
                )
            ],
            missing_mandatory=["German"],
            missing_preferred=[],
            warning="Role requires Fluent / C1 German (job posting written in German), which is not listed in your profile.",
        ),
    )
    assert assessment.fit_score == 92
    assert assessment.language_match is not None
    assert assessment.language_match.is_matched is False
    assert "German" in assessment.language_match.missing_mandatory
    assert "German" in assessment.language_match.warning


@pytest.mark.asyncio
async def test_assess_job_posting_forwards_authoritative_candidate_profile(
    db_session: AsyncSession,
):
    from langchain_core.runnables import RunnableLambda

    from app.schemas.llm import JobAssessmentResult
    from app.services.llm import assess_job_posting

    mock_assessment = JobAssessmentResult(
        position="Principal Distributed Systems Architect",
        company="HighScale",
        fit_score=95,
        programmatic_match_score=90,
        matching_skills=["Python", "Go", "Distributed Systems"],
        missing_skills=[],
        summary="Outstanding match with verified 4.0 years focused experience.",
    )

    captured_inputs = {}

    async def mock_chain_run(inputs):
        nonlocal captured_inputs
        captured_inputs = inputs
        return mock_assessment

    with patch("app.services.llm.get_task_chat_model") as mock_get_chat:
        mock_llm = MagicMock()
        mock_llm.with_structured_output.return_value = RunnableLambda(mock_chain_run)
        mock_get_chat.return_value = mock_llm

        res = await assess_job_posting(
            db_session,
            job_description="HighScale is hiring a Principal Distributed Systems Architect...",
            candidate_skills=["Python", "Go", "Distributed Systems"],
            candidate_cv="Alex Morgan - 8 years casual tenure / maintenance",
            candidate_domain_breakdown="Distributed Systems (4.0 yrs)",
            candidate_spoken_languages="English (Native)",
            candidate_years_of_experience=4.0,
            programmatic_baseline=90,
        )

        assert res.fit_score == 95
        prompt_text = captured_inputs.to_string()
        assert "Total Verified Professional Experience: 4.0 years" in prompt_text
        assert "Python, Go, Distributed Systems" in prompt_text
        assert "Distributed Systems (4.0 yrs)" in prompt_text
        assert "Authoritative Candidate Profile Priority" in prompt_text
        assert "Technical Bar Raiser" in prompt_text
        assert "Critical Risks" in prompt_text


def test_calibrate_assessment_score_and_recommendation_bounds():
    from app.services.llm import calibrate_assessment_score_and_recommendation

    # 1. Baseline 40%, raw AI 95% -> Clamped to max 40 + 25 = 65% (STRETCH_ROLE)
    score, rec = calibrate_assessment_score_and_recommendation(
        raw_fit_score=95, programmatic_baseline=40, critical_risks=[]
    )
    assert score == 65
    assert rec == "STRETCH_ROLE"

    # 2. Baseline 80%, raw AI 95%, 0 risks -> 95% (APPLY_STRONGLY)
    score, rec = calibrate_assessment_score_and_recommendation(
        raw_fit_score=95, programmatic_baseline=80, critical_risks=[]
    )
    assert score == 95
    assert rec == "APPLY_STRONGLY"

    # 3. Baseline 80%, raw AI 95%, 2 critical risks -> 95%, but downgraded to STRETCH_ROLE
    score, rec = calibrate_assessment_score_and_recommendation(
        raw_fit_score=95,
        programmatic_baseline=80,
        critical_risks=["Missing Kubernetes in production", "Seniority gap"],
    )
    assert score == 95
    assert rec == "STRETCH_ROLE"

    # 4. Baseline 80%, raw AI 30% -> Clamped to min 80 - 25 = 55% (STRETCH_ROLE)
    score, rec = calibrate_assessment_score_and_recommendation(
        raw_fit_score=30, programmatic_baseline=80, critical_risks=[]
    )
    assert score == 55
    assert rec == "STRETCH_ROLE"

    # 5. Baseline None (0 JD skills), raw AI 90% -> Capped at 70% (APPLY_MODERATELY)
    score, rec = calibrate_assessment_score_and_recommendation(
        raw_fit_score=90, programmatic_baseline=None, critical_risks=[]
    )
    assert score == 70
    assert rec == "APPLY_MODERATELY"

    # 6. Baseline 90%, raw AI 95%, but UNDERQUALIFIED seniority -> Capped at 65% (STRETCH_ROLE)
    score, rec = calibrate_assessment_score_and_recommendation(
        raw_fit_score=95,
        programmatic_baseline=90,
        critical_risks=[],
        seniority_fit="UNDERQUALIFIED",
    )
    assert score == 65
    assert rec == "STRETCH_ROLE"

    # 7. Low score < 50% -> DO_NOT_APPLY
    score, rec = calibrate_assessment_score_and_recommendation(
        raw_fit_score=35,
        programmatic_baseline=30,
        critical_risks=["Missing core stack"],
    )
    assert score == 35
    assert rec == "DO_NOT_APPLY"


@pytest.mark.asyncio
async def test_cv_anonymization_prompt_template_variables(db_session):
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.runnables import RunnableLambda

    from app.core.prompts import DEFAULT_PROMPTS, get_prompt_template
    from app.schemas.candidate_profile import CVAnonymizationResult
    from app.services.llm import anonymize_and_parse_cv

    template_str = await get_prompt_template(db_session, "cv_anonymization")
    prompt = ChatPromptTemplate.from_template(template_str)
    assert prompt.input_variables == ["resume_text"]

    # Also test DEFAULT_PROMPTS directly
    default_prompt = ChatPromptTemplate.from_template(
        DEFAULT_PROMPTS["cv_anonymization"]
    )
    assert default_prompt.input_variables == ["resume_text"]

    # Verify chain invocation with mock model
    mock_parsed = CVAnonymizationResult(
        anonymized_resume="[Candidate] Software Engineer",
        extracted_skills=["Python", "PostgreSQL"],
        total_years_experience=5.0,
        domain_expertise=["Backend"],
        domain_breakdown=[],
        spoken_languages=[{"language": "English", "proficiency": "Native"}],
        summary="Senior Backend Engineer",
    )

    with patch("app.services.llm.get_task_chat_model") as mock_get_chat:
        mock_llm = MagicMock()
        mock_llm.with_structured_output.return_value = RunnableLambda(
            lambda x: mock_parsed
        )
        mock_get_chat.return_value = mock_llm

        res = await anonymize_and_parse_cv(
            db_session, "Sample resume content text here for testing"
        )
        assert res.total_years_experience == 5.0
        assert "Python" in res.extracted_skills
