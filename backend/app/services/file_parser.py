import email
import hashlib
import io
import logging
import uuid
from datetime import UTC, datetime
from email import policy
from email.utils import parsedate_to_datetime

from app.schemas.intake import EmailPayload

logger = logging.getLogger(__name__)


def _extract_ics_summary(ics_bytes: bytes) -> str:
    """Extracts summary and date info from raw .ics calendar payload."""
    try:
        text = ics_bytes.decode("utf-8", errors="replace")
        lines = text.splitlines()
        extracted: list[str] = []
        for line in lines:
            if line.startswith("SUMMARY:"):
                extracted.append(f"Meeting Title: {line[8:].strip()}")
            elif line.startswith("DTSTART"):
                extracted.append(f"Start: {line.split(':', 1)[-1].strip()}")
            elif line.startswith("DTEND"):
                extracted.append(f"End: {line.split(':', 1)[-1].strip()}")
            elif line.startswith("DESCRIPTION:"):
                extracted.append(f"Description: {line[12:].strip()}")
        if extracted:
            return "\n[Calendar Event (.ics)]\n" + "\n".join(extracted)
    except Exception as err:
        logger.warning("Failed parsing .ics attachment: %s", err)
    return ""


def parse_eml(content: bytes) -> EmailPayload:
    """Parses raw RFC 822 / MIME .eml bytes into EmailPayload."""
    msg = email.message_from_bytes(content, policy=policy.default)

    subject = msg.get("Subject", "No Subject").strip()
    msg.get("From", "").strip()
    message_id = msg.get("Message-ID", f"eml-{uuid.uuid4().hex[:12]}").strip()
    conversation_id = msg.get("Thread-Topic") or msg.get("In-Reply-To") or message_id

    # Parse date
    received_at = datetime.now(UTC)
    date_header = msg.get("Date")
    if date_header:
        try:
            parsed_dt = parsedate_to_datetime(date_header)
            if parsed_dt.tzinfo is None:
                parsed_dt = parsed_dt.replace(tzinfo=UTC)
            received_at = parsed_dt
        except Exception:
            pass

    # Extract body and calendar attachments
    body_parts: list[str] = []
    ics_parts: list[str] = []

    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            content_disposition = str(part.get("Content-Disposition", ""))

            if (
                content_type == "text/plain"
                and "attachment" not in content_disposition
                or (
                    content_type == "text/html"
                    and not body_parts
                    and "attachment" not in content_disposition
                )
            ):
                payload = part.get_payload(decode=True)
                if payload:
                    body_parts.append(
                        payload.decode(
                            part.get_content_charset() or "utf-8", errors="replace"
                        )
                    )
            elif (
                content_type == "text/calendar"
                or part.get_filename()
                and part.get_filename().endswith(".ics")
            ):
                raw_ics = part.get_payload(decode=True)
                if raw_ics:
                    ics_summary = _extract_ics_summary(raw_ics)
                    if ics_summary:
                        ics_parts.append(ics_summary)
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            body_parts.append(
                payload.decode(msg.get_content_charset() or "utf-8", errors="replace")
            )

    full_body = "\n".join(body_parts).strip() if body_parts else msg.get_payload() or ""
    if isinstance(full_body, list):
        full_body = "\n".join(str(p) for p in full_body)
    if ics_parts:
        full_body += "\n" + "\n".join(ics_parts)

    return EmailPayload(
        conversation_id=conversation_id,
        message_id=message_id,
        received_at=received_at,
        subject=subject or "No Subject",
        body=full_body,
    )


def parse_msg(content: bytes) -> EmailPayload:
    """Parses Microsoft Outlook .msg binary bytes into EmailPayload."""
    import extract_msg

    msg = extract_msg.Message(io.BytesIO(content))
    subject = msg.subject or "No Subject"
    message_id = msg.messageId or f"msg-{uuid.uuid4().hex[:12]}"
    conversation_id = msg.conversationTopic or message_id

    received_at = msg.date or datetime.now(UTC)
    if received_at.tzinfo is None:
        received_at = received_at.replace(tzinfo=UTC)

    body = msg.body or msg.htmlBody or ""
    if isinstance(body, bytes):
        body = body.decode("utf-8", errors="replace")

    ics_parts: list[str] = []
    if hasattr(msg, "attachments"):
        for att in msg.attachments:
            if (
                hasattr(att, "longFilename")
                and att.longFilename
                and att.longFilename.endswith(".ics")
            ):
                ics_summary = _extract_ics_summary(att.data)
                if ics_summary:
                    ics_parts.append(ics_summary)

    if ics_parts:
        body += "\n" + "\n".join(ics_parts)

    return EmailPayload(
        conversation_id=conversation_id,
        message_id=message_id,
        received_at=received_at,
        subject=subject,
        body=body,
    )


def parse_txt(content: bytes, filename: str = "upload.txt") -> EmailPayload:
    """Parses plaintext / raw thread text into EmailPayload."""
    text = content.decode("utf-8", errors="replace").strip()
    lines = text.splitlines()

    subject = ""
    body_lines = []
    in_headers = True

    # Simple check for basic header block
    for line in lines:
        if in_headers:
            lower = line.lower()
            if lower.startswith("subject:"):
                subject = line.split(":", 1)[1].strip()
            elif (
                lower.startswith("from:")
                or lower.startswith("date:")
                or lower.startswith("to:")
            ):
                continue
            elif line.strip() == "":
                in_headers = False
            else:
                in_headers = False
                body_lines.append(line)
        else:
            body_lines.append(line)

    body = "\n".join(body_lines).strip() if body_lines else text
    if not subject:
        # Take first non-empty line up to 80 chars or fallback to filename
        first_line = lines[0].strip() if lines else filename
        subject = first_line[:80] if len(first_line) > 0 else filename

    # Compute deterministic message_id from content hash
    content_hash = hashlib.sha256(content).hexdigest()[:16]
    msg_id = f"txt-{content_hash}"

    return EmailPayload(
        conversation_id=f"conv-{content_hash}",
        message_id=msg_id,
        received_at=datetime.now(UTC),
        subject=subject,
        body=body,
    )


def parse_uploaded_file(filename: str, content: bytes) -> EmailPayload:
    """Dispatches file parsing based on extension (.eml, .msg, .txt)."""
    fn_lower = filename.lower()
    if fn_lower.endswith(".eml"):
        return parse_eml(content)
    elif fn_lower.endswith(".msg"):
        return parse_msg(content)
    elif fn_lower.endswith(".txt"):
        return parse_txt(content, filename=filename)
    else:
        # Default to text parser for other formats (e.g. .log, .raw)
        return parse_txt(content, filename=filename)


def normalize_resume_text(text: str) -> str:
    """Intelligently cleans and normalizes extracted CV/resume text while preserving document structure.

    Heals fragmented lines (orphaned single-word line breaks from PDF coordinates),
    normalizes excessive whitespace/blank lines, aligns bullet points onto individual lines,
    and removes spaces before punctuation.
    """
    if not text:
        return ""

    import re

    # 1. Normalize line endings and non-breaking spaces
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\xa0", " ").replace("\u200b", "").replace("\ufeff", "")

    # 2. Normalize bullet points in middle of text (ensure bullets start on new lines)
    for b in ("●", "•", "▪", "▫", "◆", "◦"):
        text = re.sub(rf"(?<!\n)[ \t]*{re.escape(b)}[ \t]*", f"\n{b} ", text)

    # Break embedded dash/bullet items onto new lines (e.g. "...updates. - Managed...")
    text = re.sub(r"(?<!\n)[ \t]+-[ \t]+", "\n- ", text)

    # 3. Collapse multiple horizontal spaces/tabs into a single space
    text = re.sub(r"[^\S\n]+", " ", text)

    # 4. Clean up spaces before punctuation (horizontal space only, preserving newlines for .NET etc.)
    text = re.sub(r"[ \t]+([,.:;!?])", r"\1", text)
    text = re.sub(r"\([ \t]+", "(", text)
    text = re.sub(r"[ \t]+\)", ")", text)
    text = re.sub(r"(\w+)[ \t]+-[ \t]*based", r"\1-based", text, flags=re.IGNORECASE)

    # 5. Process lines: heal broken/fragmented lines while preserving sections & bullets
    raw_lines = [line.strip() for line in text.split("\n")]
    merged_lines: list[str] = []

    bullet_prefixes = ("●", "•", "▪", "▫", "◆", "◦", "-", "*", "–", "—")
    known_headers = {
        "Professional Summary",
        "Technical Skills",
        "Professional Experience",
        "Work Experience",
        "Experience",
        "Education",
        "Certifications",
        "Projects",
        "Summary",
        "Core Competencies",
        "Languages",
    }
    connecting_words = (
        "and",
        "or",
        "to",
        "via",
        "in",
        "for",
        "with",
        "into",
        "using",
        "by",
        "from",
        "including",
        "through",
        "of",
        "custom",
        "local",
        "structured",
        "vision-capable",
        "deepen",
    )

    for line in raw_lines:
        if not line:
            if merged_lines and merged_lines[-1] != "":
                merged_lines.append("")
            continue

        if not merged_lines:
            merged_lines.append(line)
            continue

        # Get previous non-empty line content if any
        if merged_lines[-1] == "":
            prev_line = merged_lines[-2] if len(merged_lines) >= 2 else ""
            has_blank_separator = True
        else:
            prev_line = merged_lines[-1]
            has_blank_separator = False

        is_current_bullet = any(line.startswith(b) for b in bullet_prefixes) or bool(
            re.match(r"^\d+[\.\)]\s+", line)
        )

        is_prev_header = prev_line.endswith(":") or (
            any(prev_line.startswith(f"{b} ") for b in bullet_prefixes)
            and len(prev_line.split()) <= 4
            and not prev_line.endswith((".", ","))
        )

        is_section_header = (
            len(line.split()) <= 5
            and not line.endswith((".", ",", ";"))
            and not is_current_bullet
            and (
                line in known_headers
                or (line.isupper() and len(line) > 2)
            )
        )

        is_prev_section = (
            prev_line in known_headers
            or (len(prev_line.split()) <= 5 and prev_line.isupper() and len(prev_line) > 2)
            or is_prev_header
        )

        if is_current_bullet or is_section_header:
            merged_lines.append(line)
            continue

        if is_prev_section:
            merged_lines.append(line)
            continue

        # Check connection continuation signals
        is_prev_connector = (
            prev_line.endswith((",", ";", "-", "/", "(", "[", "&"))
            or any(prev_line.lower().endswith(f" {w}") for w in connecting_words)
            or prev_line.lower() in connecting_words
        )

        is_prev_sentence_end = prev_line.endswith((".", "!", "?"))
        is_short_fragment = (
            len(line.split()) <= 2
            or len(line) <= 25
            or line[0].islower()
            or line.startswith((",", ";", ")", "]", "/"))
        )

        if is_prev_connector:
            should_merge = True
        elif is_prev_sentence_end:
            # If previous ended with period, only merge if no blank separator and it's a continuing fragment
            should_merge = not has_blank_separator and (
                line[0].islower() or line.startswith((",", ";", ")", "]", "/"))
            )
        else:
            # Previous did not end with period or colon -> part of paragraph/list
            should_merge = True

        if should_merge and prev_line:
            if has_blank_separator:
                # Pop the empty separator line since this was a fragmented break
                merged_lines.pop()
            merged_lines[-1] = f"{merged_lines[-1]} {line}"
        else:
            merged_lines.append(line)

    # 6. Final pass: clean extra spaces and collapse redundant empty lines
    result: list[str] = []
    for line in merged_lines:
        trimmed = line.strip()
        if not trimmed:
            if result and result[-1] != "":
                result.append("")
        else:
            cleaned_line = re.sub(r"[^\S\n]+", " ", trimmed)
            cleaned_line = re.sub(r"\s+([,.:;!?])", r"\1", cleaned_line)
            result.append(cleaned_line)

    return "\n".join(result).strip()


def parse_cv_document(filename: str, content: bytes) -> str:
    """Extracts and normalizes text content from uploaded resume documents (.pdf, .docx, .doc, .txt)."""
    fn_lower = filename.lower()
    raw_text = ""

    if fn_lower.endswith(".pdf"):
        import pypdf

        try:
            reader = pypdf.PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                extracted = ""
                # Attempt layout mode first for better positional layout fidelity
                try:
                    extracted = page.extract_text(extraction_mode="layout")
                except Exception:
                    extracted = ""

                if not extracted or not extracted.strip():
                    extracted = page.extract_text()

                if extracted:
                    text_parts.append(extracted)

            raw_text = "\n\n".join(text_parts).strip()
            if not raw_text:
                raise ValueError("Could not extract any text from PDF document.")
        except Exception as err:
            logger.error("Failed parsing PDF file '%s': %s", filename, err)
            raise ValueError(f"Failed parsing PDF document: {err!s}") from err

    elif fn_lower.endswith(".docx"):
        import docx

        try:
            doc = docx.Document(io.BytesIO(content))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
            raw_text = "\n".join(text_parts).strip()
            if not raw_text:
                raise ValueError("Could not extract any text from Word document.")
        except Exception as err:
            logger.error("Failed parsing DOCX file '%s': %s", filename, err)
            raise ValueError(f"Failed parsing Word document: {err!s}") from err

    elif fn_lower.endswith(".txt"):
        try:
            raw_text = content.decode("utf-8", errors="replace").strip()
            if not raw_text:
                raise ValueError("Uploaded text file is empty.")
        except Exception as err:
            raise ValueError(f"Failed reading text file: {err!s}") from err

    elif fn_lower.endswith(".doc"):
        try:
            printable_text = content.decode("utf-8", errors="ignore")
            printable = "".join(
                c for c in printable_text if c.isprintable() or c in ("\n", "\r", "\t")
            )
            cleaned = "\n".join(
                line.strip() for line in printable.splitlines() if len(line.strip()) > 3
            )
            if cleaned and len(cleaned) > 20:
                raw_text = cleaned
            else:
                raise ValueError(
                    "Legacy .doc format text extraction yielded empty or unreadable content. Please convert file to .docx or .pdf format."
                )
        except Exception as err:
            raise ValueError(
                "Legacy .doc format requires conversion to .docx or .pdf."
            ) from err

    else:
        raise ValueError(
            f"Unsupported file type: {filename}. Supported formats are .pdf, .docx, .doc, and .txt."
        )

    # Run structure-preserving normalization on extracted text
    return normalize_resume_text(raw_text)
